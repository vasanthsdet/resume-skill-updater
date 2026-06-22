"""
Updates a .docx resume in-place (copy → edit → save).

Sections rewritten:
  1. Professional Summary  — all bullet paragraphs (same format as original)
  2. Technical Skills      — all category lines
  3. Job 1 bullets         — first entry under Professional Experience (Toyota)
  4. Job 2 bullets         — second entry (Equifax)

Everything else (name, contact, dates, company headers, education) is untouched.

Returns (output_path, change_log) where change_log is a list of dicts:
  {"section": str, "bullet": int, "skills_added": [str], "text": str}
"""

import shutil
import os
import re
import json
import anthropic
from docx import Document
from dotenv import load_dotenv

load_dotenv()

SUMMARY_KEYWORDS    = {"professional summary", "summary", "profile", "objective"}
SKILLS_KEYWORDS     = {"technical skills", "skills", "core competencies", "expertise"}
EXPERIENCE_KEYWORDS = {"professional experience", "work experience", "employment history"}


def _client() -> anthropic.Anthropic:
    key = os.getenv("ANTHROPIC_API_KEY")
    if not key:
        raise RuntimeError("ANTHROPIC_API_KEY not set — add it to .env")
    return anthropic.Anthropic(api_key=key)


# ---------------------------------------------------------------------------
# Document structure helpers
# ---------------------------------------------------------------------------

def _is_heading(para) -> bool:
    """
    Detect section headings like "Professional Summary:" while NOT
    mistaking skill-category lines like "Cloud: AWS, GCP, Azure" for headings.
    """
    text = para.text.strip()
    if not text:
        return False
    if "heading" in para.style.name.lower():
        return True
    # Pure section label — ends with colon, nothing substantive after it
    if text.rstrip().endswith(":") and len(text) < 60:
        return True
    # Bold short line — but skip "Category: value, value, ..." lines
    has_bold = para.runs and all(r.bold for r in para.runs if r.text.strip())
    if has_bold and len(text) < 80:
        colon_idx = text.find(":")
        if 0 < colon_idx < len(text) - 2:
            return False
        return True
    return False


def _find_section(doc: Document, keywords: set) -> tuple[int, list[int]]:
    paras = doc.paragraphs
    for i, para in enumerate(paras):
        lower = para.text.strip().lower()
        if any(kw in lower for kw in keywords) and _is_heading(para):
            content: list[int] = []
            j = i + 1
            while j < len(paras):
                if _is_heading(paras[j]) and paras[j].text.strip():
                    break
                if paras[j].text.strip():
                    content.append(j)
                j += 1
            return i, content
    return -1, []


def _find_job_blocks(doc: Document, exp_heading_idx: int) -> list[tuple[list[int], list[int]]]:
    paras = doc.paragraphs
    jobs: list[tuple[list[int], list[int]]] = []
    current_headers: list[int] = []
    current_bullets: list[int] = []
    in_job = False
    i = exp_heading_idx + 1

    while i < len(paras):
        text = paras[i].text.strip()
        if not text:
            i += 1
            continue

        lower = text.lower()
        if lower.rstrip().endswith(":") and any(
            kw in lower for kw in {"education", "certification", "award", "publication"}
        ):
            break

        is_company = " | " in text and bool(re.search(r"\d{4}", text))

        if is_company:
            if in_job:
                jobs.append((current_headers[:], current_bullets[:]))
            current_headers = [i]
            current_bullets = []
            in_job = True
            i += 1
            if i < len(paras):
                nxt = paras[i].text.strip()
                if nxt and len(nxt) < 70 and " | " not in nxt and not re.search(r"\d{4}", nxt):
                    current_headers.append(i)
                    i += 1
        elif in_job:
            current_bullets.append(i)
            i += 1
        else:
            i += 1

    if in_job:
        jobs.append((current_headers[:], current_bullets[:]))

    return jobs


def _get_resume_text(path: str) -> str:
    doc = Document(path)
    return "\n".join(p.text for p in doc.paragraphs if p.text.strip())


# ---------------------------------------------------------------------------
# Change detection
# ---------------------------------------------------------------------------

def _detect_changes(
    section: str,
    original: list[str],
    updated: list[str],
    skills: list[str],
) -> list[dict]:
    """
    Compare original vs updated bullets. For each changed line, record
    which recruiter skills now appear in the new text but not the old.
    """
    changes = []
    for i, (old, new) in enumerate(zip(original, updated)):
        if old.strip() == new.strip():
            continue
        added = [
            s for s in skills
            if s.lower() in new.lower() and s.lower() not in old.lower()
        ]
        changes.append({
            "section": section,
            "bullet": i + 1,
            "skills_added": added,
            "text": new,
        })
    return changes


# ---------------------------------------------------------------------------
# AI rewrite
# ---------------------------------------------------------------------------

def _call_ai(
    resume_text: str,
    recruiter_skills: list[str],
    current_summary: list[str],
    current_skills_lines: list[str],
    job1_context: str,
    job1_bullets: list[str],
    job2_context: str,
    job2_bullets: list[str],
) -> dict:

    skills_block = "\n".join(f"- {s}" for s in recruiter_skills)

    # Show the first 3 summary bullets as the format reference so the AI
    # produces bullets in exactly the same style (length, voice, structure).
    format_examples = "\n".join(f"  → {b}" for b in current_summary[:3])

    def numbered(lst: list[str]) -> str:
        return "\n".join(f"{j + 1}. {b}" for j, b in enumerate(lst))

    prompt = f"""You are an expert resume writer updating a senior QA Engineer's resume.

RECRUITER'S REQUIRED SKILLS:
{skills_block}

════════════════════════════════════════════════
WRITING RULES — follow every one strictly:

1. HUMAN VOICE: Write like a real engineer describing their own work.
   No buzzwords. No passive voice. No "responsible for" filler.

2. INTEGRATE naturally: Blend the recruiter skills INTO real accomplishments.
   ✅ Good: "Replaced the existing JMeter suite with k6, cutting CI run time by 35%."
   ❌ Bad:  "Used k6 and JMeter for performance testing responsibilities."

3. NEVER FABRICATE: Only add a skill where it genuinely fits the existing work context.

4. SAME COUNT: Return EXACTLY the same number of items as each current section.

5. SUMMARY FORMAT — this is critical:
   The Professional Summary must always follow this exact style from the original:
{format_examples}
   Each bullet must be ONE complete sentence that:
   - Opens with a noun phrase OR strong past-tense verb describing a real capability
   - Includes a specific detail (tool, metric, domain, or outcome)
   - Reads as the candidate would actually describe themselves — not a job listing

6. SKILLS FORMAT: Keep the SAME category labels. Only update the values within each line.

7. EXPERIENCE BULLETS: Maintain the same confident, first-person-implied narrative tone.
════════════════════════════════════════════════

SECTION 1 — PROFESSIONAL SUMMARY  (need exactly {len(current_summary)} bullets):
{numbered(current_summary)}

SECTION 2 — TECHNICAL SKILLS  (need exactly {len(current_skills_lines)} lines):
{chr(10).join(current_skills_lines)}

SECTION 3 — JOB 1 BULLETS  (need exactly {len(job1_bullets)} bullets)
Role: {job1_context}
{numbered(job1_bullets)}

SECTION 4 — JOB 2 BULLETS  (need exactly {len(job2_bullets)} bullets)
Role: {job2_context}
{numbered(job2_bullets)}

Full resume (context only — do not copy verbatim):
{resume_text[:2500]}

Return ONLY valid JSON — no markdown, no extra keys:
{{
  "summary_bullets": [ /* exactly {len(current_summary)} strings */ ],
  "skills_lines":    [ /* exactly {len(current_skills_lines)} strings */ ],
  "job1_bullets":    [ /* exactly {len(job1_bullets)} strings */ ],
  "job2_bullets":    [ /* exactly {len(job2_bullets)} strings */ ]
}}"""

    resp = _client().messages.create(
        model="claude-sonnet-4-6",
        max_tokens=6000,
        messages=[{"role": "user", "content": prompt}],
    )
    raw = resp.content[0].text.strip()
    raw = re.sub(r"^```(?:json)?\n?", "", raw)
    raw = re.sub(r"\n?```$", "", raw)
    return json.loads(raw)


# ---------------------------------------------------------------------------
# Document writing
# ---------------------------------------------------------------------------

def _set_para_text(para, new_text: str):
    """Replace paragraph text, preserving the first run's character formatting."""
    if not para.runs:
        para.add_run(new_text)
        return
    r = para.runs[0]
    saved = {
        "name":   r.font.name,
        "size":   r.font.size,
        "bold":   r.bold,
        "italic": r.italic,
    }
    for run in para.runs:
        run.text = ""
    r.text      = new_text
    r.font.name = saved["name"]
    r.font.size = saved["size"]
    r.bold      = saved["bold"]
    r.italic    = saved["italic"]


def _apply(doc: Document, idxs: list[int], new_texts: list[str]):
    for j, idx in enumerate(idxs):
        _set_para_text(doc.paragraphs[idx], new_texts[j] if j < len(new_texts) else "")


# ---------------------------------------------------------------------------
# Public entry point
# ---------------------------------------------------------------------------

def update_resume(
    base_path: str,
    recruiter_skills: list[str],
    output_path: str,
) -> tuple[str, list[dict]]:
    """
    Copy base_path → output_path, then rewrite Summary / Skills / Job1 / Job2.
    Returns (output_path, change_log).
    change_log is a list of {"section", "bullet", "skills_added", "text"} dicts.
    """
    shutil.copy2(base_path, output_path)
    doc = Document(output_path)
    resume_text = _get_resume_text(base_path)

    _, summary_idxs = _find_section(doc, SUMMARY_KEYWORDS)
    _, skills_idxs  = _find_section(doc, SKILLS_KEYWORDS)
    exp_idx, _       = _find_section(doc, EXPERIENCE_KEYWORDS)

    orig_summary      = [doc.paragraphs[i].text.strip() for i in summary_idxs]
    orig_skills_lines = [doc.paragraphs[i].text.strip() for i in skills_idxs]

    jobs = _find_job_blocks(doc, exp_idx) if exp_idx >= 0 else []

    job1_context, orig_job1, job1_bullet_idxs = "", [], []
    job2_context, orig_job2, job2_bullet_idxs = "", [], []

    if len(jobs) >= 1:
        h, b = jobs[0]
        job1_context     = doc.paragraphs[h[0]].text.strip()
        job1_bullet_idxs = b
        orig_job1        = [doc.paragraphs[i].text.strip() for i in b]

    if len(jobs) >= 2:
        h, b = jobs[1]
        job2_context     = doc.paragraphs[h[0]].text.strip()
        job2_bullet_idxs = b
        orig_job2        = [doc.paragraphs[i].text.strip() for i in b]

    print(f"  summary={len(orig_summary)} | skills={len(orig_skills_lines)} | "
          f"job1={len(orig_job1)} bullets | job2={len(orig_job2)} bullets")
    print("  Calling Claude Sonnet to rewrite sections...")

    ai = _call_ai(
        resume_text, recruiter_skills,
        orig_summary, orig_skills_lines,
        job1_context, orig_job1,
        job2_context, orig_job2,
    )

    new_summary = ai.get("summary_bullets", [])
    new_skills  = ai.get("skills_lines",    [])
    new_job1    = ai.get("job1_bullets",    [])
    new_job2    = ai.get("job2_bullets",    [])

    _apply(doc, summary_idxs,     new_summary)
    _apply(doc, skills_idxs,      new_skills)
    _apply(doc, job1_bullet_idxs, new_job1)
    _apply(doc, job2_bullet_idxs, new_job2)

    doc.save(output_path)

    # Build change log — compare original vs updated for each section
    change_log: list[dict] = []
    change_log += _detect_changes("Professional Summary", orig_summary,      new_summary, recruiter_skills)
    change_log += _detect_changes("Technical Skills",    orig_skills_lines,  new_skills,  recruiter_skills)
    change_log += _detect_changes(job1_context,          orig_job1,          new_job1,    recruiter_skills)
    change_log += _detect_changes(job2_context,          orig_job2,          new_job2,    recruiter_skills)

    return output_path, change_log
