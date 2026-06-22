"""
Post-process the rewritten resume to apply bold formatting:
  1. Skills section  : category label bold, values normal font
  2. Summary bullets : technology / tool names bold inline
  3. All job bullets : technology / tool names bold inline

Usage:
  python bold_skills.py
  python bold_skills.py --input resume/Revathi_Battina_Resume.docx --output resume/Revathi_Battina_Resume.docx
"""

import re, shutil, argparse, copy
from docx import Document
from lxml import etree

NS    = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
XML_NS = "http://www.w3.org/XML/1998/namespace"

SUMMARY_KEYWORDS    = {"professional summary", "summary", "profile", "objective"}
SKILLS_KEYWORDS     = {"technical skills", "skills", "core competencies", "expertise"}
EXPERIENCE_KEYWORDS = {"professional experience", "work experience", "employment history"}

EXTRA_TERMS = [
    "Playwright", "Cypress", "Selenium WebDriver", "Selenium", "Appium",
    "WebDriverIO", "Perfecto", "REST Assured", "Karate", "Postman", "Bruno",
    "Insomnia", "JMeter", "LoadRunner", "gRPC", "SOAP", "SoapUI", "Axios",
    "Charles Proxy", "TestNG", "JUnit", "Cucumber", "Gherkin", "Spock",
    "Mocha", "Gauge", "Robot Framework", "Jenkins", "GitHub Actions",
    "GitLab CI", "Azure DevOps", "Maven", "Gradle", "Docker", "Kubernetes",
    "AWS Lambda", "API Gateway", "DynamoDB", "CodeCommit", "CloudWatch",
    "AWS", "GCP", "Azure", "Git", "GitHub", "GitLab", "JIRA", "X-Ray",
    "SonarQube", "Confluence", "IntelliJ", "VS Code", "Dynatrace", "Splunk",
    "ELK Stack", "Prometheus", "Grafana", "Salesforce", "Espresso", "XCUITest",
    "Xcode", "Android Studio", "GitHub Copilot", "ChatGPT", "Claude",
    "Java", "Python", "TypeScript", "JavaScript", "Ruby", "Groovy", "Go",
    "NodeJS", "NPM", "C#", ".NET", "Oracle", "PostgreSQL", "MySQL",
    "MongoDB", "Couchbase", "SQL", "NoSQL", "Linux",
    "CI/CD", "BDD", "DDD", "TDD", "REST", "Swagger", "OAuth", "JWT",
    "HL7", "EDI", "Agile", "Scrum", "S3",
]


# ── Document structure helpers ─────────────────────────────────────────────────

def _is_heading(para) -> bool:
    text = para.text.strip()
    if not text:
        return False
    if "heading" in para.style.name.lower():
        return True
    if text.rstrip().endswith(":") and len(text) < 60:
        return True
    has_bold = para.runs and all(r.bold for r in para.runs if r.text.strip())
    if has_bold and len(text) < 80:
        colon_idx = text.find(":")
        if 0 < colon_idx < len(text) - 2:
            return False
        return True
    return False


def _find_section(doc, keywords):
    paras = doc.paragraphs
    for i, para in enumerate(paras):
        lower = para.text.strip().lower()
        if any(kw in lower for kw in keywords) and _is_heading(para):
            content = []
            j = i + 1
            while j < len(paras):
                if _is_heading(paras[j]) and paras[j].text.strip():
                    break
                if paras[j].text.strip():
                    content.append(j)
                j += 1
            return i, content
    return -1, []


def _find_all_job_blocks(doc, exp_heading_idx):
    paras = doc.paragraphs
    jobs = []
    label, bullets, in_job = "", [], False
    i = exp_heading_idx + 1
    while i < len(paras):
        text = paras[i].text.strip()
        if not text:
            i += 1
            continue
        lower = text.lower()
        if lower.rstrip().endswith(":") and any(
            kw in lower for kw in {"education", "certification", "award"}
        ):
            break
        if " | " in text:
            if in_job:
                jobs.append((label, bullets[:]))
            label, bullets, in_job = text, [], True
            i += 1
            if i < len(paras):
                nxt = paras[i].text.strip()
                if nxt and len(nxt) < 70 and " | " not in nxt and not re.search(r"\d{4}", nxt):
                    i += 1
        elif in_job:
            bullets.append(i)
            i += 1
        else:
            i += 1
    if in_job:
        jobs.append((label, bullets[:]))
    return jobs


# ── Low-level XML run helpers ──────────────────────────────────────────────────

def _set_run_bold_xml(r_elem, is_bold: bool):
    """Set <w:b> and <w:bCs> directly on a <w:r> XML element."""
    rpr = r_elem.find(f"{{{NS}}}rPr")
    if rpr is None:
        rpr = etree.Element(f"{{{NS}}}rPr")
        r_elem.insert(0, rpr)

    for tag in (f"{{{NS}}}b", f"{{{NS}}}bCs"):
        for old in rpr.findall(tag):
            rpr.remove(old)

    b   = etree.Element(f"{{{NS}}}b")
    bcs = etree.Element(f"{{{NS}}}bCs")
    if not is_bold:
        b.set(f"{{{NS}}}val", "0")
        bcs.set(f"{{{NS}}}val", "0")
    rpr.insert(0, bcs)
    rpr.insert(0, b)


def _remove_ppr_bold(para):
    """Strip bold from paragraph-level default run properties so runs control it."""
    ppr = para._element.find(f"{{{NS}}}pPr")
    if ppr is None:
        return
    rpr = ppr.find(f"{{{NS}}}rPr")
    if rpr is None:
        return
    for tag in (f"{{{NS}}}b", f"{{{NS}}}bCs"):
        for elem in rpr.findall(tag):
            rpr.remove(elem)


def _rewrite_runs(para, segments: list[tuple[str, bool]]):
    """
    Replace every <w:r> in the paragraph with one run per segment.

    Strategy: deepcopy the first existing run (preserves font/size/colour/etc.),
    then update only the bold flag and text.  Insert runs at the exact position
    where the original first run sat so ordering stays correct.
    """
    existing = para._element.findall(f"{{{NS}}}r")
    if not existing:
        # Fallback: no existing runs — use add_run
        for text, is_bold in segments:
            if text:
                r = para.add_run(text)
                r.bold = is_bold
        return

    ref    = existing[0]
    p_elem = para._element
    children = list(p_elem)
    insert_at = children.index(ref)

    # Remove all existing runs
    for r in existing:
        p_elem.remove(r)

    # Insert new runs in their place
    offset = 0
    for text, is_bold in segments:
        if not text:
            continue

        new_r = copy.deepcopy(ref)

        # Update bold
        _set_run_bold_xml(new_r, is_bold)

        # Replace text nodes
        for t_elem in new_r.findall(f"{{{NS}}}t"):
            new_r.remove(t_elem)
        t_elem = etree.Element(f"{{{NS}}}t")
        t_elem.text = text
        if text != text.strip():           # has leading/trailing spaces
            t_elem.set(f"{{{XML_NS}}}space", "preserve")
        new_r.append(t_elem)

        p_elem.insert(insert_at + offset, new_r)
        offset += 1


# ── Skill-term pattern ─────────────────────────────────────────────────────────

def _build_pattern(skill_terms: list[str]) -> re.Pattern:
    sorted_terms = sorted({t for t in skill_terms if t}, key=len, reverse=True)
    parts = []
    for t in sorted_terms:
        pre  = r"\b" if re.match(r"\w", t[0])  else ""
        post = r"\b" if re.match(r"\w", t[-1]) else ""
        parts.append(pre + re.escape(t) + post)
    return re.compile("|".join(parts), re.IGNORECASE) if parts else re.compile(r"(?!)")


def _split_segments(text: str, pattern: re.Pattern) -> list[tuple[str, bool]]:
    segs, pos = [], 0
    for m in pattern.finditer(text):
        if m.start() > pos:
            segs.append((text[pos:m.start()], False))
        segs.append((m.group(), True))
        pos = m.end()
    if pos < len(text):
        segs.append((text[pos:], False))
    return segs


# ── Public formatting functions ────────────────────────────────────────────────

def format_skills_line(para):
    """Skills section: category key bold, values normal."""
    text = para.text
    c = text.find(":")
    if c < 0:
        return
    _remove_ppr_bold(para)
    segments = [(text[:c + 1], True)]
    if text[c + 1:]:
        segments.append((text[c + 1:], False))
    _rewrite_runs(para, segments)


def bold_skills_in_bullet(para, pattern: re.Pattern):
    """Inline-bold every matched skill/tool in a bullet paragraph."""
    text = para.text
    if not text.strip():
        return
    segs = _split_segments(text, pattern)
    if not any(bold for _, bold in segs):
        return
    _rewrite_runs(para, segs)


# ── Orchestrator ───────────────────────────────────────────────────────────────

def apply_bold_formatting(input_path: str, output_path: str):
    if input_path != output_path:
        shutil.copy2(input_path, output_path)
    doc = Document(output_path)

    _, skills_idxs  = _find_section(doc, SKILLS_KEYWORDS)
    _, summary_idxs = _find_section(doc, SUMMARY_KEYWORDS)
    exp_idx, _       = _find_section(doc, EXPERIENCE_KEYWORDS)
    jobs = _find_all_job_blocks(doc, exp_idx) if exp_idx >= 0 else []

    print(f"Detected  summary={len(summary_idxs)} bullets  "
          f"skills={len(skills_idxs)} lines  "
          f"jobs={len(jobs)}")

    # 1. Skills section
    print(f"\nSkills: key bold, values normal ({len(skills_idxs)} lines)...")
    for idx in skills_idxs:
        format_skills_line(doc.paragraphs[idx])

    # 2. Build term list from skills section
    skill_terms = list(EXTRA_TERMS)
    for idx in skills_idxs:
        line = doc.paragraphs[idx].text
        c = line.find(":")
        if c >= 0:
            for term in re.split(r",\s*", line[c + 1:].strip()):
                t = term.strip().strip("()")
                if t and len(t) > 1:
                    skill_terms.append(t)
    pattern = _build_pattern(skill_terms)

    # 3. Summary
    print(f"Summary: bolding skill terms in {len(summary_idxs)} bullets...")
    for idx in summary_idxs:
        bold_skills_in_bullet(doc.paragraphs[idx], pattern)

    # 4. Job bullets
    total = sum(len(b) for _, b in jobs)
    print(f"Experience: bolding skill terms in {total} bullets across {len(jobs)} jobs...")
    for _, bullet_idxs in jobs:
        for idx in bullet_idxs:
            bold_skills_in_bullet(doc.paragraphs[idx], pattern)

    doc.save(output_path)
    print(f"\nSaved -> {output_path}")


def main():
    parser = argparse.ArgumentParser(description="Bold skill/tech terms across resume sections")
    parser.add_argument("--input",  default="resume/Revathi_Battina_Resume.docx")
    parser.add_argument("--output", default="resume/Revathi_Battina_Resume.docx")
    args = parser.parse_args()
    print(f"Reading: {args.input}\n")
    apply_bold_formatting(args.input, args.output)
    print("Done. Open in Word to verify bold formatting.")


if __name__ == "__main__":
    main()
