"""
Scans every job section in the resume and removes duplicate bullet points.

Two-pass strategy:
  Pass 1 — Jaccard (fast, no API):
    Process jobs in order (Toyota -> Equifax -> UST -> Availity -> Sana Technos).
    Keep the FIRST occurrence; remove exact and near-duplicate (>82% word overlap)
    copies in later jobs.

  Pass 2 — AI within-job (Claude Haiku, --ai flag):
    For each job independently, ask Claude to identify bullets that:
      - Cover the same topic as another bullet in the same job
      - Reference outdated/obsolete tools
      - Are too generic for the role
    Removes the weaker duplicate, keeping the more specific bullet.

Usage:
  python deduplicate.py                   # Jaccard pass only, reads/overwrites resume/base_resume.docx
  python deduplicate.py --ai              # Jaccard + AI pass (recommended for within-job clean-up)
  python deduplicate.py --dry-run --ai    # Preview both passes, no file written
  python deduplicate.py --input in.docx --output out.docx --ai
"""

import os
import re
import json
import argparse
import httpx
import anthropic
from docx import Document
from dotenv import load_dotenv

load_dotenv()

EXPERIENCE_KEYWORDS   = {"professional experience", "work experience", "employment history"}
CROSS_JOB_THRESHOLD   = 0.82   # Jaccard threshold for cross-job exact/near-duplicates
WITHIN_JOB_THRESHOLD  = 0.50   # Jaccard threshold for within-job near-duplicates


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _normalize(text: str) -> str:
    return re.sub(r"\s+", " ", text.strip().lower())


def _jaccard(a: str, b: str) -> float:
    wa, wb = set(a.split()), set(b.split())
    if not wa or not wb:
        return 0.0
    return len(wa & wb) / len(wa | wb)


def _is_heading(para) -> bool:
    text = para.text.strip()
    if not text:
        return False
    if "heading" in para.style.name.lower():
        return True
    if text.rstrip().endswith(":") and len(text) < 60:
        return True
    has_bold = para.runs and all(r.bold for r in para.runs if r.text.strip())
    if has_bold and len(text) < 80:
        colon_idx = text.find(":")
        if 0 < colon_idx < len(text) - 2:
            return False
        return True
    return False


def _find_experience_heading(doc: Document) -> int:
    for i, para in enumerate(doc.paragraphs):
        lower = para.text.strip().lower()
        if any(kw in lower for kw in EXPERIENCE_KEYWORDS) and _is_heading(para):
            return i
    return -1


def _find_all_job_blocks(doc: Document, exp_heading_idx: int) -> list[tuple[str, list[int]]]:
    """Return [(job_label, [bullet_para_idxs]), ...] for every job after the experience heading."""
    paras   = doc.paragraphs
    jobs: list[tuple[str, list[int]]] = []
    label   = ""
    bullets: list[int] = []
    in_job  = False
    i = exp_heading_idx + 1

    while i < len(paras):
        text = paras[i].text.strip()
        if not text:
            i += 1
            continue

        lower = text.lower()
        if lower.rstrip().endswith(":") and any(
            kw in lower for kw in {"education", "certification", "award", "publication"}
        ):
            break

        if " | " in text:
            if in_job:
                jobs.append((label, bullets[:]))
            label   = text
            bullets = []
            in_job  = True
            i += 1
            if i < len(paras):
                nxt = paras[i].text.strip()
                if nxt and len(nxt) < 70 and " | " not in nxt and not re.search(r"\d{4}", nxt):
                    i += 1
        elif in_job:
            bullets.append(i)
            i += 1
        else:
            i += 1

    if in_job:
        jobs.append((label, bullets[:]))

    return jobs


def _remove_para(para) -> None:
    elem = para._element
    elem.getparent().remove(elem)


# ---------------------------------------------------------------------------
# Pass 1 — Jaccard cross-job deduplication
# ---------------------------------------------------------------------------

def _jaccard_pass(
    paras: list,
    jobs: list[tuple[str, list[int]]],
    dry_run: bool,
) -> list[int]:
    """
    Cross-job deduplication: keep first occurrence, remove later copies.
    Also catches within-job near-duplicates above WITHIN_JOB_THRESHOLD.
    Returns list of para indices to remove.
    """
    seen_exact: set[str]  = set()
    seen_cross: list[str] = []   # for cross-job Jaccard
    to_remove:  list[int] = []

    for job_label, bullet_idxs in jobs:
        seen_within: list[str] = []   # reset per job for within-job Jaccard
        removed_here: list[str] = []

        for idx in bullet_idxs:
            text = paras[idx].text.strip()
            if not text:
                continue
            norm = _normalize(text)

            # Exact match (cross-job or within-job)
            if norm in seen_exact:
                to_remove.append(idx)
                removed_here.append(text[:90])
                continue

            # Within-job near-duplicate (lower threshold)
            if any(_jaccard(norm, s) >= WITHIN_JOB_THRESHOLD for s in seen_within):
                to_remove.append(idx)
                removed_here.append(text[:90])
                continue

            # Cross-job near-duplicate (higher threshold)
            if any(_jaccard(norm, s) >= CROSS_JOB_THRESHOLD for s in seen_cross):
                to_remove.append(idx)
                removed_here.append(text[:90])
                continue

            seen_exact.add(norm)
            seen_within.append(norm)
            seen_cross.append(norm)

        if removed_here:
            short = job_label.split("|")[0].strip()[:30]
            print(f"  [Jaccard] {short}: {len(removed_here)} removed")
            for t in removed_here:
                print(f"    - {t}...")

    return to_remove


# ---------------------------------------------------------------------------
# Pass 2 — AI within-job deduplication
# ---------------------------------------------------------------------------

def _ai_dedup_job(job_label: str, bullets: list[str]) -> list[int]:
    """
    Ask Claude Haiku to identify redundant bullet indices within one job.
    Returns 0-based indices of bullets to remove.
    """
    key = os.getenv("ANTHROPIC_API_KEY")
    if not key:
        raise RuntimeError("ANTHROPIC_API_KEY not set — add it to .env")
    client = anthropic.Anthropic(api_key=key, http_client=httpx.Client(verify=False))

    numbered = "\n".join(f"{i}. {b}" for i, b in enumerate(bullets))

    prompt = f"""You are reviewing bullet points for this resume job section: {job_label}

Identify bullets to REMOVE because they:
1. Cover the same topic as another bullet in this list — keep the more specific/detailed one, remove the vaguer one
2. Reference outdated or obsolete tools (e.g. Firebug was discontinued in 2017 and should not appear in a 2024-2025 role)
3. Are too generic for a Senior QA/SDET Engineer (e.g. "Experienced with POM pattern", "Followed Agile-Scrum process")

BULLETS:
{numbered}

Rules:
- Only remove clear duplicates or obvious problems. When in doubt, keep the bullet.
- For topic duplicates, prefer the bullet with more detail and concrete outcomes.
- Return ONLY a valid JSON array of 0-based indices to remove. Example: [2, 7, 12]
- If nothing should be removed, return: []"""

    resp = client.messages.create(
        model="claude-haiku-4-5-20251001",
        max_tokens=300,
        messages=[{"role": "user", "content": prompt}],
    )
    raw = resp.content[0].text.strip()
    # Extract the first JSON array found anywhere in the response
    match = re.search(r"\[[\d,\s]*\]", raw)
    if not match:
        return []
    result = json.loads(match.group())
    return [int(x) for x in result if isinstance(x, (int, float))]


def _ai_pass(
    paras: list,
    jobs: list[tuple[str, list[int]]],
    already_removed: set[int],
    dry_run: bool,
) -> list[int]:
    """
    For each job, use AI to find topic-based duplicates the Jaccard pass missed.
    Returns additional para indices to remove.
    """
    to_remove: list[int] = []

    for job_label, bullet_idxs in jobs:
        # Work only on bullets that survived the Jaccard pass
        live_idxs   = [i for i in bullet_idxs if i not in already_removed]
        live_bullets = [paras[i].text.strip() for i in live_idxs if paras[i].text.strip()]
        live_idxs    = [i for i in live_idxs if paras[i].text.strip()]

        if len(live_bullets) < 2:
            continue

        short = job_label.split("|")[0].strip()[:30]
        print(f"  [AI] Scanning {short} ({len(live_bullets)} bullets)...")

        try:
            remove_positions = _ai_dedup_job(job_label, live_bullets)
        except Exception as e:
            print(f"  [AI] Skipped {short} due to error: {e}")
            continue

        removed_here: list[str] = []
        for pos in remove_positions:
            if 0 <= pos < len(live_idxs):
                para_idx = live_idxs[pos]
                to_remove.append(para_idx)
                removed_here.append(live_bullets[pos][:90])

        if removed_here:
            print(f"  [AI] {short}: {len(removed_here)} removed")
            for t in removed_here:
                print(f"    - {t}...")
        else:
            print(f"  [AI] {short}: nothing to remove")

    return to_remove


# ---------------------------------------------------------------------------
# Main entry point
# ---------------------------------------------------------------------------

def deduplicate(
    input_path: str,
    output_path: str,
    dry_run: bool = False,
    use_ai: bool  = False,
) -> int:
    """
    Remove duplicate bullets. Returns total count removed.
    """
    doc   = Document(input_path)
    paras = doc.paragraphs

    exp_idx = _find_experience_heading(doc)
    if exp_idx < 0:
        print("  [dedup] Experience section not found.")
        if not dry_run:
            doc.save(output_path)
        return 0

    jobs = _find_all_job_blocks(doc, exp_idx)
    print(f"  Found {len(jobs)} job block(s)\n")

    # Pass 1 — Jaccard
    print("Pass 1 — Jaccard (exact + similarity):")
    jaccard_remove = _jaccard_pass(paras, jobs, dry_run)

    # Pass 2 — AI
    ai_remove: list[int] = []
    if use_ai:
        print("\nPass 2 — AI (within-job topic analysis):")
        ai_remove = _ai_pass(paras, jobs, set(jaccard_remove), dry_run)

    all_remove = sorted(set(jaccard_remove + ai_remove))
    total = len(all_remove)

    if not total:
        print("\nNo duplicates found.")
        if not dry_run:
            doc.save(output_path)
        return 0

    if dry_run:
        print(f"\n[dry-run] Would remove {total} bullet(s) total. No file written.")
        return total

    paras_to_remove = [doc.paragraphs[i] for i in all_remove]
    for para in paras_to_remove:
        _remove_para(para)

    doc.save(output_path)
    print(f"\nDone — {total} duplicate bullet(s) removed.")
    print(f"Saved -> {output_path}")
    return total


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------

def main():
    parser = argparse.ArgumentParser(description="Remove duplicate resume bullet points")
    parser.add_argument("--input",   default="resume/base_resume.docx")
    parser.add_argument("--output",  default="resume/base_resume.docx")
    parser.add_argument("--dry-run", action="store_true", help="Preview only, no file change")
    parser.add_argument("--ai",      action="store_true", help="Use AI for within-job topic dedup (recommended)")
    args = parser.parse_args()

    print(f"Reading: {args.input}\n")
    removed = deduplicate(args.input, args.output, dry_run=args.dry_run, use_ai=args.ai)
    if not args.dry_run:
        print(f"Total removed: {removed}")


if __name__ == "__main__":
    main()
