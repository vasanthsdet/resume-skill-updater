"""
Full resume rewrite — human voice, 4-page target.

Target layout:
  Professional Summary  ->  15 bullets
  Technical Skills      ->  12 category lines (same labels)
  Toyota (current)      ->  16 bullets
  Equifax               ->  15 bullets
  UST Global            ->  12 bullets
  Availity              ->   9 bullets
  Sana Technos          ->   5 bullets

Usage:
  python rewrite_resume.py
  python rewrite_resume.py --input resume/base_resume.docx --output resume/Revathi_Battina_Resume.docx
"""

import copy, os, re, json, shutil, httpx, argparse
import anthropic
from docx import Document
from dotenv import load_dotenv
from bold_skills import apply_bold_formatting

load_dotenv()

TARGETS = {
    "summary": 15,
    "skills":  12,
    "jobs": [16, 15, 12, 9, 5],   # Toyota, Equifax, UST, Availity, Sana Technos
}

SUMMARY_KEYWORDS    = {"professional summary", "summary", "profile", "objective"}
SKILLS_KEYWORDS     = {"technical skills", "skills", "core competencies", "expertise"}
EXPERIENCE_KEYWORDS = {"professional experience", "work experience", "employment history"}


def _client() -> anthropic.Anthropic:
    key = os.getenv("ANTHROPIC_API_KEY")
    if not key:
        raise RuntimeError("ANTHROPIC_API_KEY not set in .env")
    return anthropic.Anthropic(api_key=key, http_client=httpx.Client(verify=False))


# ── Document structure helpers ─────────────────────────────────────────────────

def _is_heading(para) -> bool:
    text = para.text.strip()
    if not text:
        return False
    if "heading" in para.style.name.lower():
        return True
    if text.rstrip().endswith(":") and len(text) < 60:
        return True
    has_bold = para.runs and all(r.bold for r in para.runs if r.text.strip())
    if has_bold and len(text) < 80:
        colon_idx = text.find(":")
        if 0 < colon_idx < len(text) - 2:
            return False
        return True
    return False


def _find_section(doc: Document, keywords: set) -> tuple[int, list[int]]:
    paras = doc.paragraphs
    for i, para in enumerate(paras):
        lower = para.text.strip().lower()
        if any(kw in lower for kw in keywords) and _is_heading(para):
            content: list[int] = []
            j = i + 1
            while j < len(paras):
                if _is_heading(paras[j]) and paras[j].text.strip():
                    break
                if paras[j].text.strip():
                    content.append(j)
                j += 1
            return i, content
    return -1, []


def _find_all_job_blocks(doc: Document, exp_heading_idx: int) -> list[tuple[str, list[int]]]:
    paras = doc.paragraphs
    jobs: list[tuple[str, list[int]]] = []
    label, bullets, in_job = "", [], False
    i = exp_heading_idx + 1
    while i < len(paras):
        text = paras[i].text.strip()
        if not text:
            i += 1
            continue
        lower = text.lower()
        if lower.rstrip().endswith(":") and any(kw in lower for kw in {"education", "certification", "award"}):
            break
        if " | " in text:
            if in_job:
                jobs.append((label, bullets[:]))
            label, bullets, in_job = text, [], True
            i += 1
            if i < len(paras):
                nxt = paras[i].text.strip()
                if nxt and len(nxt) < 70 and " | " not in nxt and not re.search(r"\d{4}", nxt):
                    i += 1
        elif in_job:
            bullets.append(i)
            i += 1
        else:
            i += 1
    if in_job:
        jobs.append((label, bullets[:]))
    return jobs


# ── Document writing ───────────────────────────────────────────────────────────

def _set_para_text(para, new_text: str):
    if not para.runs:
        para.add_run(new_text)
        return
    r = para.runs[0]
    saved = {"name": r.font.name, "size": r.font.size, "bold": r.bold, "italic": r.italic}
    for run in para.runs:
        run.text = ""
    r.text = new_text; r.font.name = saved["name"]; r.font.size = saved["size"]
    r.bold = saved["bold"]; r.italic = saved["italic"]


def _clone_insert_after(ref_para, new_text: str):
    NS     = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    new_el = copy.deepcopy(ref_para._element)
    ref_para._element.addnext(new_el)
    t_elems = new_el.findall(f".//{{{NS}}}t")
    if t_elems:
        t_elems[0].text = new_text
        for t in t_elems[1:]:
            t.text = ""


def _apply_section(doc: Document, idxs: list[int], new_texts: list[str]):
    """Write new_texts into idxs paragraphs; removes surplus or clones if needed."""
    if not idxs:
        return
    paras = doc.paragraphs
    for j in range(min(len(idxs), len(new_texts))):
        _set_para_text(paras[idxs[j]], new_texts[j])
    if len(idxs) > len(new_texts):
        surplus = [paras[i] for i in idxs[len(new_texts):]]
        for p in surplus:
            p._element.getparent().remove(p._element)
    elif len(new_texts) > len(idxs):
        ref = doc.paragraphs[idxs[-1]]
        for extra in new_texts[len(idxs):]:
            _clone_insert_after(ref, extra)
            all_elems = [p._element for p in doc.paragraphs]
            ref = doc.paragraphs[all_elems.index(ref._element) + 1]


# ── AI rewrite ─────────────────────────────────────────────────────────────────

def _call_ai(summary, skills, jobs_input) -> dict:
    jobs_block = ""
    for k, (label, bullets) in enumerate(jobs_input):
        t = TARGETS["jobs"][k] if k < len(TARGETS["jobs"]) else len(bullets)
        jobs_block += f"\nJOB {k+1} — {label}  (need exactly {t} bullets)\n"
        jobs_block += "\n".join(f"  {j+1}. {b}" for j, b in enumerate(bullets)) + "\n"

    job_count_lines = "\n".join(
        f"   Job {k+1} ({lbl.split(chr(124))[0].strip()[:22]}) -> {TARGETS['jobs'][k] if k < len(TARGETS['jobs']) else '?'} bullets"
        for k, (lbl, _) in enumerate(jobs_input)
    )

    prompt = f"""You are a professional resume writer. Rewrite this entire QA Engineer resume
so every bullet sounds like the candidate wrote it herself — natural, direct, specific.
The result must fill 4 printed pages.

Candidate: Revathi Battina — Sr. QA Automation Engineer, 11+ years experience, at Toyota.

════════════════════════════════════════════
WRITING RULES:

1. HUMAN VOICE — write how a senior engineer actually talks about her work.
   No passive voice. No "responsible for", "involved in", "assisted with", "worked on".
   Every sentence should feel like she said it in a panel interview.

2. CONCRETE — name the actual tool, the actual problem, ideally a result or metric.
   Bad:  "Improved test coverage."
   Good: "Grew API test coverage from 40% to 92% by writing REST Assured contract tests
          for every microservice and embedding them in the Jenkins pipeline."

3. VARIED OPENERS — mix action verbs (Built, Designed, Led, Drove, Shipped, Owned,
   Replaced, Introduced, Took) with noun-led sentences ("The Playwright suite I built...",
   "Our CI pipeline..."). Never repeat the same opener three bullets in a row.

4. BULLET LENGTH — 28-42 words each. Rich, specific, fills a line and a half.
   Two tight sentences are fine. No fragments, no run-ons.

5. EXACT COUNTS — return these precisely:
   Summary        ->  {TARGETS['summary']} bullets
   Technical Skills -> {TARGETS['skills']} lines (same category labels, reorder tools by relevance)
{job_count_lines}

6. NO DUPLICATES — each bullet must cover a distinct accomplishment.

7. TONE BY SENIORITY:
   Toyota / Equifax  : architectural decisions, metrics, leadership, cloud-native work.
   UST / Availity    : frameworks built, team breadth, delivery impact.
   Sana Technos      : foundation, learning curve, core testing skills.

8. SUMMARY — 15 single-sentence bullets. Each opens with a noun phrase or strong verb.
   Covers: years + domain, key tools, leadership, cloud, AI-in-QA, performance, API,
   mobile, what sets her apart. Do NOT start any bullet with "I".

9. SKILLS LINES — keep exact category labels (Languages:, UI Automation:, etc.).
   Reorder tools within each line to put the most in-demand first.
════════════════════════════════════════════

CURRENT CONTENT:

PROFESSIONAL SUMMARY  ({len(summary)} bullets -> rewrite as {TARGETS['summary']}):
{chr(10).join(f'{j+1}. {b}' for j, b in enumerate(summary))}

TECHNICAL SKILLS  ({len(skills)} lines -> keep {TARGETS['skills']}):
{chr(10).join(skills)}
{jobs_block}

Return ONLY valid JSON — no markdown, no extra text:
{{
  "summary": [ /* exactly {TARGETS['summary']} strings */ ],
  "skills":  [ /* exactly {TARGETS['skills']} strings */ ],
  "jobs": [
    {', '.join(f'[/* {TARGETS["jobs"][k] if k < len(TARGETS["jobs"]) else "?"} strings */]' for k in range(len(jobs_input)))}
  ]
}}"""

    resp = _client().messages.create(
        model="claude-sonnet-4-6",
        max_tokens=8000,
        messages=[{"role": "user", "content": prompt}],
    )
    raw = resp.content[0].text.strip()
    raw = re.sub(r"^```(?:json)?\n?", "", raw)
    raw = re.sub(r"\n?```$", "", raw)
    return json.loads(raw)


# ── Main ───────────────────────────────────────────────────────────────────────

def rewrite_resume(input_path: str, output_path: str) -> str:
    shutil.copy2(input_path, output_path)
    doc = Document(output_path)

    _, summary_idxs = _find_section(doc, SUMMARY_KEYWORDS)
    _, skills_idxs  = _find_section(doc, SKILLS_KEYWORDS)
    exp_idx, _       = _find_section(doc, EXPERIENCE_KEYWORDS)

    current_summary = [doc.paragraphs[i].text.strip() for i in summary_idxs]
    current_skills  = [doc.paragraphs[i].text.strip() for i in skills_idxs]
    raw_jobs        = _find_all_job_blocks(doc, exp_idx) if exp_idx >= 0 else []
    jobs_input      = [(lbl, [doc.paragraphs[i].text.strip() for i in idxs]) for lbl, idxs in raw_jobs]

    print("Current -> Target:")
    print(f"  Summary        : {len(current_summary):>3} -> {TARGETS['summary']}")
    print(f"  Skills         : {len(current_skills):>3} -> {TARGETS['skills']}")
    for k, (lbl, idxs) in enumerate(raw_jobs):
        t = TARGETS["jobs"][k] if k < len(TARGETS["jobs"]) else len(idxs)
        print(f"  {lbl.split('|')[0].strip()[:20]:<20}: {len(idxs):>3} -> {t}")

    print("\nCalling Claude Sonnet for full rewrite (~45s)...")
    ai = _call_ai(current_summary, current_skills, jobs_input)

    # Apply bottom-to-top so removing/inserting paragraphs in one section
    # does not shift the stored indices of sections above it.
    for k in range(len(raw_jobs) - 1, -1, -1):
        _, idxs = raw_jobs[k]
        if k < len(ai.get("jobs", [])):
            _apply_section(doc, idxs, ai["jobs"][k])
    _apply_section(doc, skills_idxs,  ai["skills"])
    _apply_section(doc, summary_idxs, ai["summary"])

    doc.save(output_path)
    print(f"Saved -> {output_path}")

    print("Applying bold formatting...")
    apply_bold_formatting(output_path, output_path)
    return output_path


def main():
    parser = argparse.ArgumentParser(description="Full resume rewrite — human voice, 4-page target")
    parser.add_argument("--input",  default="resume/base_resume.docx")
    parser.add_argument("--output", default="resume/Revathi_Battina_Resume.docx")
    args = parser.parse_args()
    print(f"Reading: {args.input}\n")
    rewrite_resume(args.input, args.output)
    print("\nDone. Open in Word to verify page count.")


if __name__ == "__main__":
    main()
